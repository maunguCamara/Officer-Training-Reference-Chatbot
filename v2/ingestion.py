import os
import re
import json
import shutil
from pathlib import Path
from dotenv import load_dotenv
load_dotenv()

import fitz
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from langdetect import detect, DetectorFactory
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma

DetectorFactory.seed = 0

PDF_DIR = Path("data/pdfs")
CHROMA_DB_DIR = Path("data/chroma_db")

# Embeddings selection (same logic as main.py)
LLM_PROVIDER = os.getenv("LLM_PROVIDER", "openai").lower()
USE_E5 = os.getenv("USE_E5", "false").lower() == "true"

if LLM_PROVIDER == "ollama":
    if USE_E5:
        from sentence_transformers import SentenceTransformer
        class E5Embeddings:
            def __init__(self):
                self.model = SentenceTransformer("intfloat/multilingual-e5-large")
            def embed_documents(self, texts):
                return self.model.encode(["passage: " + t for t in texts], normalize_embeddings=True)
            def embed_query(self, text):
                return self.model.encode("query: " + text, normalize_embeddings=True)
        embeddings = E5Embeddings()
    else:
        from langchain_huggingface import HuggingFaceEmbeddings
        embeddings = HuggingFaceEmbeddings(
            model_name="BAAI/bge-m3",
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True}
        )
else:
    from langchain_openai import OpenAIEmbeddings
    embeddings = OpenAIEmbeddings(model="text-embedding-3-large")

# ========== Extraction Functions ==========
def extract_text_from_pdf(file_path):
    doc = fitz.open(file_path)
    docs = []
    for page_num in range(len(doc)):
        text = doc[page_num].get_text("text").strip()
        if text:
            text = " ".join(text.split())
            docs.append(Document(page_content=text, metadata={"source": file_path.name, "page": page_num+1}))
    doc.close()
    return docs

def extract_text_from_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    # Create one document per paragraph (simulating pages)
    return [Document(page_content="\n".join(full_text), metadata={"source": file_path.name, "page": 1})]

def extract_text_from_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    return [Document(page_content=text, metadata={"source": file_path.name, "page": 1})]

def extract_text_from_html(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f.read(), 'html.parser')
    text = soup.get_text()
    return [Document(page_content=text, metadata={"source": file_path.name, "page": 1})]

def extract_documents(doc_dir):
    all_docs = []
    for file_path in doc_dir.iterdir():
        if file_path.suffix == '.pdf':
            all_docs.extend(extract_text_from_pdf(file_path))
        elif file_path.suffix == '.docx':
            all_docs.extend(extract_text_from_docx(file_path))
        elif file_path.suffix == '.txt':
            all_docs.extend(extract_text_from_txt(file_path))
        elif file_path.suffix == '.html':
            all_docs.extend(extract_text_from_html(file_path))
    return all_docs

def split_documents(docs, chunk_size=500, chunk_overlap=50):
    splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
        chunk_size=chunk_size, chunk_overlap=chunk_overlap
    )
    return splitter.split_documents(docs)

def add_language_metadata(chunks):
    for chunk in chunks:
        try:
            lang = detect(chunk.page_content)
            if lang.startswith("en"): chunk.metadata["language"] = "en"
            elif lang.startswith("sw"): chunk.metadata["language"] = "sw"
            else: chunk.metadata["language"] = "en"
        except:
            chunk.metadata["language"] = "unknown"
    return chunks

# --- Topic Extraction ---
def roman_to_int(s):
    s = s.upper()
    rom_val = {'I':1, 'V':5, 'X':10, 'L':50, 'C':100, 'D':500, 'M':1000}
    total = 0
    prev = 0
    for ch in reversed(s):
        val = rom_val.get(ch, 0)
        if val >= prev:
            total += val
            prev = val
        else:
            total -= val
    return total

def extract_part_topics_from_pdf(file_path):
    doc = fitz.open(file_path)
    parts = []
    seen = set()
    for page_num in range(len(doc)):
        text = doc[page_num].get_text("text")
        lines = text.split('\n')
        for line in lines:
            m = re.match(r'PART\s+([IVXLCDM]+)\s*[—–-]\s*(.*)', line.strip(), re.IGNORECASE)
            if m:
                roman = m.group(1)
                title = m.group(2).strip()
                part_num = roman_to_int(roman)
                # Clean title
                clean_title = re.split(r'\s+\d+\s*[—–-]', title)[0].strip()
                if not clean_title:
                    clean_title = title
                if (page_num+1, roman) not in seen:
                    parts.append({"id": str(part_num), "title": f"Part {roman}: {clean_title}", "page": page_num+1})
                    seen.add((page_num+1, roman))
    doc.close()
    parts.sort(key=lambda x: x["page"])
    for i, p in enumerate(parts, start=1):
        p["id"] = str(i)
    return parts

def extract_headings_from_pdf(file_path):
    doc = fitz.open(file_path)
    headings = []
    seen = set()
    for page_num in range(len(doc)):
        text = doc[page_num].get_text("text")
        matches = re.findall(r"^(\d+\.\s+[A-Za-z][^\n]{2,})", text, re.MULTILINE)
        for match in matches:
            title = match.strip()
            if title not in seen:
                headings.append((page_num+1, title))
                seen.add(title)
    doc.close()
    headings.sort(key=lambda x: (x[0], int(re.match(r"(\d+)", x[1]).group(1)) if re.match(r"(\d+)", x[1]) else 0))
    return [{"id": str(i+1), "title": title, "page": page} for i, (page, title) in enumerate(headings)]

def extract_headings_from_docx(file_path):
    doc = DocxDocument(file_path)
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            headings.append({"title": para.text.strip(), "page": 1})
    return headings

def build_topics_json():
    topics = {}
    for file_path in PDF_DIR.iterdir():
        if file_path.suffix == '.pdf':
            parts = extract_part_topics_from_pdf(file_path)
            if not parts:
                parts = extract_headings_from_pdf(file_path)
            topics[file_path.name] = parts if parts else [{"id": "1", "title": file_path.stem, "page": 1}]
        elif file_path.suffix == '.docx':
            parts = extract_headings_from_docx(file_path)
            topics[file_path.name] = parts if parts else [{"id": "1", "title": file_path.stem, "page": 1}]
        else:
            topics[file_path.name] = [{"id": "1", "title": file_path.stem, "page": 1}]
    with open(Path("data/topics.json"), "w", encoding="utf-8") as f:
        json.dump(topics, f, indent=2)
    print("Saved topics.json")

def assign_topic_metadata(docs, topics_json):
    for doc in docs:
        source = doc.metadata.get("source", "")
        book_topics = topics_json.get(source, [])
        page = doc.metadata.get("page", 1)
        assigned_topic = None
        book_topics_sorted = sorted(book_topics, key=lambda x: x.get("page", 1))
        for t in book_topics_sorted:
            if t.get("page", 1) <= page:
                assigned_topic = t
        if assigned_topic:
            doc.metadata["topic_title"] = assigned_topic["title"].strip()
            doc.metadata["topic_id"] = assigned_topic["id"]
    return docs

# ========== Main Ingestion ==========
def update_vector_store(force_reload=False):
    print("Starting document ingestion...")
    if CHROMA_DB_DIR.exists() and force_reload:
        shutil.rmtree(CHROMA_DB_DIR)
    # Rebuild topics.json
    build_topics_json()
    with open(Path("data/topics.json"), "r", encoding="utf-8") as f:
        topics_json = json.load(f)

    raw_docs = extract_documents(PDF_DIR)
    chunks = split_documents(raw_docs)
    chunks = add_language_metadata(chunks)
    chunks = assign_topic_metadata(chunks, topics_json)

    vectordb = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=str(CHROMA_DB_DIR)
    )
    
    print(f"Ingestion complete. {len(chunks)} chunks stored.")
    return vectordb